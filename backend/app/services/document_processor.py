"""
Document processor - reads and processes various file formats
Based on document_processor.py from qwen-deployment project
"""
import os
import re
import logging
from typing import List, Dict, Optional, Union
from pathlib import Path
from docx import Document
from openpyxl import load_workbook
import pdfplumber
import PyPDF2
import json

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process documents of various formats"""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
    def load_file(self, file_path: Union[str, Path]) -> str:
        """Load file content based on extension"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Файл не найден: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.txt':
                return self._load_txt(file_path)
            elif file_extension == '.md':
                return self._load_markdown(file_path)
            elif file_extension == '.py':
                return self._load_python(file_path)
            elif file_extension in ['.json']:
                return self._load_json(file_path)
            elif file_extension == '.docx':
                return self._load_docx(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                return self._load_excel(file_path)
            elif file_extension == '.pdf':
                return self._load_pdf(file_path)
            else:
                logger.warning(f"Неизвестный формат файла {file_extension}, пробуем как текст")
                return self._load_txt(file_path)
                
        except Exception as e:
            logger.error(f"Ошибка при загрузке файла {file_path}: {e}")
            raise
    
    def _load_txt(self, file_path: Path) -> str:
        """Load text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            for encoding in ['cp1251', 'latin1', 'ascii']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            raise UnicodeDecodeError(f"Не удалось определить кодировку файла {file_path}")
    
    def _load_markdown(self, file_path: Path) -> str:
        """Load markdown file"""
        text = self._load_txt(file_path)
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text) 
        text = re.sub(r'`(.*?)`', r'\1', text)
        text = re.sub(r'```[\s\S]*?```', '', text)
        return text
    
    def _load_python(self, file_path: Path) -> str:
        """Load Python file"""
        text = self._load_txt(file_path)
        enhanced_text = f"Это Python код из файла {file_path.name}:\n\n{text}"
        
        functions = re.findall(r'^def\s+(\w+)', text, re.MULTILINE)
        classes = re.findall(r'^class\s+(\w+)', text, re.MULTILINE)
        
        if functions or classes:
            enhanced_text += f"\n\nСтруктура файла:\n"
            if classes:
                enhanced_text += f"Классы: {', '.join(classes)}\n"
            if functions:
                enhanced_text += f"Функции: {', '.join(functions)}\n"
        
        return enhanced_text
    
    def _load_json(self, file_path: Path) -> str:
        """Load JSON file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        text = f"JSON файл {file_path.name}:\n\n"
        text += json.dumps(data, ensure_ascii=False, indent=2)
        return text
    
    def _load_docx(self, file_path: Path) -> str:
        """Load Word document"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            text_parts.append(f"Word документ: {file_path.name}\n")
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("Библиотека python-docx не установлена. Установите: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Ошибка при чтении Word документа: {e}")
            raise
    
    def _load_excel(self, file_path: Path) -> str:
        """Load Excel file"""
        try:
            workbook = load_workbook(file_path, read_only=True, data_only=True)
            text_parts = []
            
            text_parts.append(f"Excel файл: {file_path.name}\n")
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Лист: {sheet_name}")
                
                rows_data = []
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell.strip() for cell in row_data):
                        rows_data.append(" | ".join(row_data))
                
                if rows_data:
                    text_parts.extend(rows_data)
                    text_parts.append("") 
            
            workbook.close()
            return "\n".join(text_parts)
            
        except ImportError:
            logger.error("Библиотека openpyxl не установлена. Установите: pip install openpyxl")
            raise
        except Exception as e:
            logger.error(f"Ошибка при чтении Excel файла: {e}")
            raise
    
    def _load_pdf(self, file_path: Path) -> str:
        """Load PDF file"""
        try:            
            text_parts = []
            text_parts.append(f"PDF документ: {file_path.name}\n")
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(f"Страница {page_num}:")
                        text_parts.append(text.strip())
                        text_parts.append("") 
            
            return "\n".join(text_parts)
            
        except ImportError:
            try:
                text_parts = []
                text_parts.append(f"PDF документ: {file_path.name}\n")
                
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        text = page.extract_text()
                        if text and text.strip():
                            text_parts.append(f"Страница {page_num}:")
                            text_parts.append(text.strip())
                            text_parts.append("")
                
                return "\n".join(text_parts)
                
            except ImportError:
                logger.error("Библиотеки для работы с PDF не установлены. Установите: pip install pdfplumber PyPDF2")
                raise
        except Exception as e:
            logger.error(f"Ошибка при чтении PDF файла: {e}")
            raise
    
    def chunk_text(self, text: str, metadata: Optional[Dict] = None) -> List[Dict[str, any]]:
        """Split text into chunks"""
        if not text.strip():
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            if end < len(text):
                for break_char in ['\n\n', '\n', '. ', '! ', '? ']:
                    break_pos = text.rfind(break_char, start, end)
                    if break_pos != -1:
                        end = break_pos + len(break_char)
                        break
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunk = {
                    'text': chunk_text,
                    'start_pos': start,
                    'end_pos': end,
                    'chunk_id': len(chunks),
                    'metadata': metadata or {}
                }
                chunks.append(chunk)
            start = max(start + 1, end - self.chunk_overlap)
        
        return chunks


